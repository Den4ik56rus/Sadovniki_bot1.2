# src/services/documents/docx_parser.py

"""
Улучшенный парсер DOCX документов.

Извлекает текст как сплошной поток без разбиения на страницы,
с сохранением информации о структуре (заголовки, списки).

Особенности:
- Сплошной текст без page breaks
- Детекция заголовков по стилям и форматированию
- Сохранение границ списков для защиты при chunking
"""

from typing import Dict, List, Optional, Any
from pathlib import Path


def extract_text_from_docx_v2(file_path: str) -> Dict[str, Any]:
    """
    Извлекает текст из DOCX как сплошной поток с информацией о структуре.

    Параметры:
        file_path: Путь к DOCX файлу

    Возвращает:
        {
            "full_text": str,           # Сплошной текст
            "headings": [               # Найденные заголовки
                {"text": "Глава 1", "position": 0, "level": 1}
            ],
            "lists": [                  # Границы списков
                {"start": 100, "end": 250, "type": "numbered"}
            ],
            "paragraphs_count": int,    # Количество параграфов
            "error": Optional[str]
        }
    """
    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        return {
            "full_text": "",
            "headings": [],
            "lists": [],
            "paragraphs_count": 0,
            "error": "python-docx library not installed. Install with: pip install python-docx"
        }

    try:
        doc = DocxDocument(file_path)

        text_parts = []
        headings = []
        lists = []
        current_position = 0

        # Состояние для отслеживания списков
        in_list = False
        list_start = 0
        list_type = None

        for para in doc.paragraphs:
            para_text = para.text.strip()

            if not para_text:
                # Пустой параграф может означать конец списка
                if in_list:
                    lists.append({
                        "start": list_start,
                        "end": current_position,
                        "type": list_type,
                    })
                    in_list = False
                continue

            # Определяем тип параграфа
            is_heading, heading_level = _detect_heading(para)
            is_list_item, item_type = _detect_list_item(para_text, para)

            # Обрабатываем заголовки
            if is_heading:
                if in_list:
                    lists.append({
                        "start": list_start,
                        "end": current_position,
                        "type": list_type,
                    })
                    in_list = False

                headings.append({
                    "text": para_text,
                    "position": current_position,
                    "level": heading_level,
                })

            # Обрабатываем списки
            if is_list_item:
                if not in_list:
                    in_list = True
                    list_start = current_position
                    list_type = item_type
            else:
                if in_list:
                    lists.append({
                        "start": list_start,
                        "end": current_position,
                        "type": list_type,
                    })
                    in_list = False

            # Добавляем текст
            text_parts.append(para_text)
            current_position += len(para_text) + 1  # +1 для \n

        # Закрываем открытый список в конце
        if in_list:
            lists.append({
                "start": list_start,
                "end": current_position,
                "type": list_type,
            })

        full_text = "\n".join(text_parts)

        return {
            "full_text": full_text,
            "headings": headings,
            "lists": lists,
            "paragraphs_count": len(text_parts),
            "error": None,
        }

    except Exception as e:
        return {
            "full_text": "",
            "headings": [],
            "lists": [],
            "paragraphs_count": 0,
            "error": str(e),
        }


def _detect_heading(paragraph) -> tuple[bool, int]:
    """
    Определяет, является ли параграф заголовком.

    Проверяет:
    1. Стиль параграфа (Heading 1, Heading 2, ...)
    2. Форматирование: жирный + короткий текст + центрирование

    Возвращает:
        (is_heading: bool, level: int)
    """
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        return False, 0

    style_name = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip()

    # 1. Проверка стиля
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
            return True, level
        except (ValueError, IndexError):
            return True, 1

    # 2. Эвристики для нестандартных заголовков
    # Жирный текст + короткая строка (до 100 символов) + центрирование
    if len(text) > 0 and len(text) <= 100:
        is_bold = _is_paragraph_bold(paragraph)
        is_centered = paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

        if is_bold and is_centered:
            return True, 2  # Считаем как H2

        # Жирный + короткий + заканчивается на двоеточие или без точки
        if is_bold and (text.endswith(":") or not text.endswith(".")):
            return True, 3  # Считаем как H3

    return False, 0


def _is_paragraph_bold(paragraph) -> bool:
    """Проверяет, весь ли параграф написан жирным."""
    if not paragraph.runs:
        return False

    for run in paragraph.runs:
        if run.text.strip():  # Только непустые runs
            if not run.bold:
                return False
    return True


def _detect_list_item(text: str, paragraph) -> tuple[bool, Optional[str]]:
    """
    Определяет, является ли параграф элементом списка.

    Проверяет:
    1. Стиль параграфа (List Paragraph, List Number, List Bullet)
    2. Текстовые паттерны (1., 2., -, *, а), б))

    Возвращает:
        (is_list_item: bool, list_type: str | None)
    """
    import re

    style_name = paragraph.style.name if paragraph.style else ""

    # 1. Проверка стиля
    if "List" in style_name:
        if "Number" in style_name or "Numbered" in style_name:
            return True, "numbered"
        elif "Bullet" in style_name:
            return True, "bulleted"
        else:
            return True, "generic"

    # 2. Текстовые паттерны
    patterns = {
        "numbered": r'^\s*\d+[\.\)]\s',       # 1. или 1)
        "bulleted": r'^\s*[-•●○◦▪]\s',        # -, •, ●, ○, ◦, ▪
        "lettered": r'^\s*[а-яa-z][\.\)]\s',  # а. или a)
        "roman": r'^\s*[ivxIVX]+[\.\)]\s',    # i., ii., I., II.
    }

    for list_type, pattern in patterns.items():
        if re.match(pattern, text):
            return True, list_type

    return False, None


def get_document_structure(file_path: str) -> Dict[str, Any]:
    """
    Возвращает структуру документа (только заголовки и метаданные).

    Полезно для предварительного анализа документа перед полной обработкой.
    """
    result = extract_text_from_docx_v2(file_path)

    return {
        "headings": result["headings"],
        "lists_count": len(result["lists"]),
        "paragraphs_count": result["paragraphs_count"],
        "total_length": len(result["full_text"]),
        "error": result["error"],
    }


def extract_text_simple(file_path: str) -> Dict[str, Any]:
    """
    Простое извлечение текста из DOCX без анализа структуры.

    Для совместимости с текущей реализацией.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return {
            "full_text": "",
            "pages": [],
            "error": "python-docx library not installed"
        }

    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        return {
            "full_text": full_text,
            "pages": [{"page_number": 1, "text": full_text}],
            "error": None
        }
    except Exception as e:
        return {
            "full_text": "",
            "pages": [],
            "error": str(e)
        }
