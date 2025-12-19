# src/services/documents/prompt_extractor.py

"""
Извлечение текста из документов для промт-системы.

Поддерживаемые форматы:
- .pages (Apple Pages) — через textutil (macOS) или QuickLook/Preview.pdf
- .docx (Microsoft Word) — через python-docx
- .pdf (PDF) — через pypdf
"""

import asyncio
import io
import logging
import os
import platform
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {'.pages', '.docx', '.pdf'}


async def extract_text_from_file(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Извлечь текст из файла документа.

    Args:
        file_path: Путь к файлу
        file_type: Тип файла ('pages', 'docx', 'pdf')

    Returns:
        {
            "text": str,       # Извлечённый текст
            "success": bool,   # Успешно ли извлечение
            "error": str|None  # Текст ошибки если failed
        }
    """
    try:
        if file_type == 'pages':
            text = await extract_from_pages(file_path)
        elif file_type == 'docx':
            text = await extract_from_docx(file_path)
        elif file_type == 'pdf':
            text = await extract_from_pdf(file_path)
        else:
            return {
                "text": "",
                "success": False,
                "error": f"Неподдерживаемый формат файла: {file_type}"
            }

        if not text or not text.strip():
            return {
                "text": "",
                "success": False,
                "error": "Не удалось извлечь текст из документа (пустой результат)"
            }

        return {
            "text": text.strip(),
            "success": True,
            "error": None
        }

    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return {
            "text": "",
            "success": False,
            "error": str(e)
        }


async def extract_from_pages(file_path: str) -> str:
    """
    Извлечь текст из .pages файла (Apple Pages).

    Методы извлечения (в порядке приоритета):
    1. AppleScript (macOS) — открывает Pages и извлекает body text
    2. textutil (macOS) — для старых версий .pages
    3. QuickLook/Preview.pdf — fallback для старых .pages файлов
    """
    # Метод 1: AppleScript через Pages.app (самый надёжный на macOS)
    if platform.system() == 'Darwin':
        try:
            text = await _extract_pages_with_applescript(file_path)
            if text and text.strip():
                return text
        except Exception as e:
            logger.warning(f"AppleScript extraction failed: {e}, trying fallback...")

    # Метод 2: textutil (только на macOS)
    if platform.system() == 'Darwin' and shutil.which('textutil'):
        try:
            text = await _extract_pages_with_textutil(file_path)
            if text and text.strip():
                return text
        except Exception as e:
            logger.warning(f"textutil extraction failed: {e}, trying fallback...")

    # Метод 3: Fallback на QuickLook/Preview.pdf
    try:
        text = await _extract_pages_from_pdf(file_path)
        if text and text.strip():
            return text
    except Exception as e:
        logger.warning(f"PDF extraction from .pages failed: {e}")

    raise ValueError(
        "Не удалось извлечь текст из .pages файла. "
        "Попробуйте экспортировать файл как .docx или .pdf в Apple Pages."
    )


async def _extract_pages_with_applescript(file_path: str) -> str:
    """Извлечь текст из .pages через AppleScript (открывает Pages.app)."""
    # Получаем абсолютный путь
    abs_path = os.path.abspath(file_path)

    # AppleScript для открытия документа, извлечения текста и закрытия
    script = f'''
    tell application "Pages"
        set theDoc to open POSIX file "{abs_path}"
        set bodyText to body text of theDoc
        close theDoc saving no
        return bodyText
    end tell
    '''

    process = await asyncio.create_subprocess_exec(
        'osascript', '-e', script,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE
    )
    stdout, stderr = await process.communicate()

    if process.returncode != 0:
        error_msg = stderr.decode('utf-8', errors='ignore') if stderr else 'Unknown error'
        raise RuntimeError(f"AppleScript failed: {error_msg}")

    return stdout.decode('utf-8', errors='ignore')


async def _extract_pages_with_textutil(file_path: str) -> str:
    """Извлечь текст из .pages через textutil (macOS)."""
    with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as tmp:
        tmp_path = tmp.name

    try:
        # textutil -convert txt -output /tmp/out.txt input.pages
        process = await asyncio.create_subprocess_exec(
            'textutil',
            '-convert', 'txt',
            '-output', tmp_path,
            file_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()

        if process.returncode != 0:
            error_msg = stderr.decode('utf-8', errors='ignore') if stderr else 'Unknown error'
            raise RuntimeError(f"textutil failed: {error_msg}")

        # Читаем результат
        with open(tmp_path, 'r', encoding='utf-8') as f:
            return f.read()

    finally:
        # Удаляем временный файл
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


async def _extract_pages_from_pdf(file_path: str) -> str:
    """Извлечь текст из QuickLook/Preview.pdf внутри .pages архива."""
    from pypdf import PdfReader

    with zipfile.ZipFile(file_path, 'r') as zf:
        namelist = zf.namelist()

        # Ищем Preview.pdf
        pdf_path = None
        for name in namelist:
            if name.endswith('Preview.pdf') or name == 'QuickLook/Preview.pdf':
                pdf_path = name
                break

        if not pdf_path:
            for name in namelist:
                if name.lower().endswith('.pdf'):
                    pdf_path = name
                    break

        if not pdf_path:
            raise ValueError("Не найден Preview.pdf в .pages файле")

        pdf_data = zf.read(pdf_path)
        reader = PdfReader(io.BytesIO(pdf_data))

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return '\n\n'.join(text_parts)


async def extract_from_docx(file_path: str) -> str:
    """
    Извлечь текст из .docx файла (Microsoft Word).
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_texts.append(cell.text.strip())
                if row_texts:
                    text_parts.append(' | '.join(row_texts))

        return '\n\n'.join(text_parts)

    except Exception as e:
        raise ValueError(f"Ошибка чтения .docx файла: {e}")


async def extract_from_pdf(file_path: str) -> str:
    """
    Извлечь текст из .pdf файла.
    """
    try:
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return '\n\n'.join(text_parts)

    except Exception as e:
        raise ValueError(f"Ошибка чтения .pdf файла: {e}")


def get_file_type(filename: str) -> Optional[str]:
    """
    Определить тип файла по расширению.

    Returns:
        'pages', 'docx', 'pdf' или None если формат не поддерживается
    """
    ext = Path(filename).suffix.lower()

    if ext == '.pages':
        return 'pages'
    elif ext == '.docx':
        return 'docx'
    elif ext == '.pdf':
        return 'pdf'

    return None


def is_supported_file(filename: str) -> bool:
    """Проверить, поддерживается ли формат файла."""
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS
